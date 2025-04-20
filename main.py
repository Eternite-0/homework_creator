import os
import cv2
import numpy as np
import random
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx2pdf import convert
from PyPDF2 import PdfReader

def rand_select(option):
    rand = random.random()
    for i in range(len(option)):
        if rand < 1 / len(option) * (i + 1):
            return option[i]

def create_dir(dir_path):
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)

def word(text):
    rand = int(random.random() * 200 - 100)
    font = ["萌妹子体", "李国夫手写体", "陈静的字完整版"]
    doc = Document()
    para = doc.add_paragraph(" ")
    for line in text:
        for i in line:
            run = para.add_run(i)
            selected_font = rand_select(font)
            run.font.name = selected_font
            run.font.size = 300000 + rand
            run.font.color.rgb = RGBColor(60, 60, 60)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), selected_font)
            rand = int(random.random() * 200 - 100)
    doc_path = os.path.join(os.path.dirname(os.path.realpath(__file__)), "temporary", "text.docx")
    doc.save(doc_path)

def word_to_jpg():
    convert(os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text.docx", os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text.pdf")
    reader = PdfReader(os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text.pdf")
    if reader.is_encrypted:
        reader.decrypt('')
    page = len(reader.pages)
    for i in range(page):
        os.system(os.path.split(os.path.realpath(__file__))[0] + "\\mutool.exe draw -o " + os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text" + str(i + 1) + ".png -w 1998 -h 2585 " + os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text.pdf " + str(i + 1))
    return page

def pho_mix(page):
    base_path = os.path.dirname(os.path.realpath(__file__))
    background_path = os.path.join(base_path, 'background.jpg')
    background = cv2.imread(background_path)
    if background is None:
        print(f"错误：无法加载背景图片: {background_path}")
        return
    offset_x = 30
    offset_y = 100

    for num in range(page):
        word_filename = f'text{num + 1}.png'
        word_path = os.path.join(base_path, 'temporary', word_filename)
        word_img = cv2.imread(word_path)
        if word_img is None:
            print(f"错误：无法加载文字图片: {word_path}")
            continue

        word_img = cv2.resize(word_img, (0, 0), fx = 1.1, fy = 1.2)

        print(f'开始处理 {num + 1}/{page}')

        gray_word = cv2.cvtColor(word_img, cv2.COLOR_BGR2GRAY)

        _, mask = cv2.threshold(gray_word, 240, 255, cv2.THRESH_BINARY_INV)

        roi_rows, roi_cols = mask.shape
        bg_rows, bg_cols, _ = background.shape

        y1, y2 = offset_y, offset_y + roi_rows
        x1, x2 = offset_x, offset_x + roi_cols

        current_roi_rows, current_roi_cols = word_img.shape[:2]
        y1, y2 = offset_y, offset_y + current_roi_rows
        x1, x2 = offset_x, offset_x + current_roi_cols

        if y1 < 0 or y2 > bg_rows or x1 < 0 or x2 > bg_cols:
             print(f"警告：第 {num+1} 页的文字区域 ({current_roi_rows}x{current_roi_cols} at {offset_x},{offset_y}) 超出背景边界 ({bg_rows}x{bg_cols})，跳过混合。")
             res = copy.deepcopy(background)
        else:
            if mask.shape[0] != current_roi_rows or mask.shape[1] != current_roi_cols:
                mask = cv2.resize(mask, (current_roi_cols, current_roi_rows), interpolation=cv2.INTER_NEAREST)

            roi = background[y1:y2, x1:x2]

            mask_inv = cv2.bitwise_not(mask)
            bg_masked = cv2.bitwise_and(roi, roi, mask=mask_inv)

            fg_masked = cv2.bitwise_and(word_img, word_img, mask=mask)

            combined_roi = cv2.add(bg_masked, fg_masked)

            res = copy.deepcopy(background)
            res[y1:y2, x1:x2] = combined_roi
        print(f'完成处理 {num + 1}')
        rows, cols, channel = res.shape
        M = cv2.getRotationMatrix2D((cols / 2, rows / 2), random.random() * 8 - 4, 1.085)
        res = cv2.warpAffine(res, M, (cols, rows))
        res_filename = f'res{num + 1}.jpg'
        res_path = os.path.join(base_path, 'res', res_filename)
        cv2.imwrite(res_path, res)

def main():
    
    path = os.path.split(os.path.realpath(__file__))[0]
    create_dir(path + '\\res')
    create_dir(path + '\\temporary')
    for i in os.listdir(path + '\\res'):
        os.remove(path + '\\res\\' + i)
    with open(path + "\\input.txt", "r", encoding='utf-8') as f:
        word(f.readlines())
    if input("continue?(Y/N)") == "Y":
        pass
    else:
        return
    pho_mix(word_to_jpg())
    for i in os.listdir(path + '\\temporary'):
        os.remove(os.path.split(os.path.realpath(__file__))[0] + '\\temporary\\'+ i)

if __name__=='__main__':
    main()